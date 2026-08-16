from __future__ import annotations

import hashlib
import io
import re
import shutil
import zipfile
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {
    ".xlsx": "xlsx",
    ".pdf": "pdf",
    ".docx": "docx",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".webp": "image",
    ".tif": "image",
    ".tiff": "image",
    ".bmp": "image",
}
IMAGE_EXTENSIONS = {extension for extension, kind in SUPPORTED_EXTENSIONS.items() if kind == "image"}
ENTERPRISE_CODE_PATTERN = re.compile(r"(?<![A-Za-z0-9])TF[A-Za-z0-9_-]{2,}(?![A-Za-z0-9])", re.IGNORECASE)
REFERENCE_ONLY_PHRASE = "转型规划结论"
LOW_CONFIDENCE_THRESHOLD = 0.80
MAX_EVIDENCE_ITEMS = 500


def classify_file(filename: str) -> str | None:
    return SUPPORTED_EXTENSIONS.get(Path(filename or "").suffix.lower())


def safe_filename(filename: str) -> str:
    base = Path(filename or "attachment").name
    base = re.sub(r"[^\w.\-\u4e00-\u9fff]+", "_", base, flags=re.UNICODE).strip("._")
    return base or "attachment"


def sha256_bytes(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _issue(code: str, severity: str, message: str, location: dict[str, Any] | None = None) -> dict[str, Any]:
    return {
        "code": code,
        "severity": severity,
        "message": message,
        "location": location,
    }


def _evidence(
    *,
    kind: str,
    location: dict[str, Any],
    text: str,
    confidence: float,
    source_field: str | None = None,
) -> dict[str, Any]:
    compact = " ".join(str(text).split())
    return {
        "evidence_id": f"ev-{len(text)}-{hashlib.sha1(compact.encode('utf-8')).hexdigest()[:12]}",
        "kind": kind,
        "location": location,
        "text_excerpt": compact[:2000],
        "confidence": round(max(0.0, min(float(confidence), 1.0)), 4),
        "source_field": source_field,
    }


def _detect_enterprises(text: str) -> list[str]:
    return sorted({match.upper() for match in ENTERPRISE_CODE_PATTERN.findall(text or "")})


def _ocr_capability() -> dict[str, Any]:
    try:
        import paddleocr  # type: ignore  # noqa: F401

        return {"available": True, "provider": "paddleocr", "reason": None}
    except Exception:
        pass
    if shutil.which("tesseract"):
        return {"available": True, "provider": "pytesseract", "reason": None}
    return {
        "available": False,
        "provider": None,
        "reason": "未检测到PaddleOCR或tesseract；扫描件/图片需要人工复核。",
    }


def _ocr_image(image: Any) -> tuple[str, float, dict[str, Any]]:
    capability = _ocr_capability()
    if not capability["available"]:
        return "", 0.0, capability
    if capability["provider"] != "pytesseract":
        return "", 0.0, {**capability, "reason": "PaddleOCR能力已探测，但本轮未启用模型初始化。"}
    try:
        import pytesseract
        from pytesseract import Output

        data = pytesseract.image_to_data(image, output_type=Output.DICT, config="--psm 6")
        words: list[str] = []
        confidences: list[float] = []
        for word, raw_conf in zip(data.get("text", []), data.get("conf", [])):
            word = str(word or "").strip()
            try:
                confidence = float(raw_conf) / 100.0
            except (TypeError, ValueError):
                confidence = 0.0
            if word:
                words.append(word)
                confidences.append(max(0.0, min(confidence, 1.0)))
        text = " ".join(words).strip()
        return text, (sum(confidences) / len(confidences) if confidences else 0.0), capability
    except Exception as exc:
        return "", 0.0, {**capability, "available": False, "reason": f"OCR执行失败：{type(exc).__name__}"}


def _finalize(
    *,
    file_type: str,
    parser: str,
    evidence: list[dict[str, Any]],
    issues: list[dict[str, Any]],
    detected_enterprise_codes: list[str],
    expected_enterprise_code: str | None,
    metadata: dict[str, Any],
) -> dict[str, Any]:
    detected = sorted({code.upper() for code in detected_enterprise_codes})
    expected = str(expected_enterprise_code or "").strip().upper() or None
    reference_only_detected = any(REFERENCE_ONLY_PHRASE in item.get("text_excerpt", "") for item in evidence)
    if reference_only_detected:
        issues.append(
            _issue(
                "reference_only_content",
                "warning",
                "材料包含“转型规划结论”字样；该内容仅可进入参考对照层，不得作为模型输入、标签或事实特征。",
            )
        )

    conflict = bool(expected and detected and (expected not in detected or len(detected) > 1))
    if conflict:
        issues.append(
            _issue(
                "enterprise_conflict",
                "error",
                f"材料识别企业代号 {', '.join(detected)} 与当前运行企业 {expected} 不一致，已阻断合并。",
            )
        )

    has_error = any(issue.get("severity") == "error" for issue in issues)
    evidence_confidence = [float(item.get("confidence", 0.0)) for item in evidence]
    average_confidence = round(sum(evidence_confidence) / len(evidence_confidence), 4) if evidence_confidence else 0.0
    low_confidence = not detected or any(value < LOW_CONFIDENCE_THRESHOLD for value in evidence_confidence)
    if not detected and evidence and not any(issue.get("code") == "ocr_unavailable" for issue in issues):
        issues.append(_issue("missing_enterprise_code", "warning", "材料有可提取内容，但未识别出企业代号，需要人工确认归属。"))
    if low_confidence and evidence and not any(issue.get("code") == "low_confidence" for issue in issues):
        issues.append(_issue("low_confidence", "warning", "至少一条候选证据低于离线解析置信度阈值，需要人工复核。"))

    if conflict:
        status = "blocked_conflict"
    elif has_error:
        status = "failed"
    elif low_confidence or reference_only_detected:
        status = "needs_review"
    else:
        status = "passed"
    merge_allowed = status == "passed" and bool(expected) and expected in detected and not reference_only_detected
    return {
        "status": status,
        "file_type": file_type,
        "parser": parser,
        "expected_enterprise_code": expected,
        "detected_enterprise_codes": detected,
        "merge_allowed": merge_allowed,
        "fact_eligible": merge_allowed,
        "confidence": average_confidence,
        "confidence_threshold": LOW_CONFIDENCE_THRESHOLD,
        "evidence": evidence[:MAX_EVIDENCE_ITEMS],
        "issues": issues,
        "metadata": metadata,
        "reference_only": {
            "detected": reference_only_detected,
            "model_context_excluded": True,
            "notice": "转型规划结论始终是参考验证层，不进入输入特征。",
        },
    }


def _parse_pdf(
    path: Path,
    expected_enterprise_code: str | None,
    external_client: Any | None = None,
) -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []
    detected: list[str] = []
    page_count = 0
    native_text_pages = 0
    table_count = 0
    parser_names: list[str] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path), strict=False)
        if reader.is_encrypted:
            return _finalize(
                file_type="pdf",
                parser="pypdf",
                evidence=[],
                issues=[_issue("encrypted_file", "error", "PDF已加密，当前离线流程不能读取。")],
                detected_enterprise_codes=[],
                expected_enterprise_code=expected_enterprise_code,
                metadata={"page_count": len(reader.pages), "ocr": _ocr_capability()},
            )
        page_count = len(reader.pages)
        for page_number, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                native_text_pages += 1
                parser_names.append("pypdf")
                evidence.append(
                    _evidence(
                        kind="pdf_text",
                        location={"page": page_number},
                        text=text,
                        confidence=0.96,
                    )
                )
                detected.extend(_detect_enterprises(text))
            else:
                issues.append(_issue("page_without_native_text", "warning", f"第{page_number}页没有可提取的原生文本。", {"page": page_number}))
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                for page_number, page in enumerate(pdf.pages, start=1):
                    for table_number, table in enumerate(page.extract_tables() or [], start=1):
                        if not table:
                            continue
                        table_count += 1
                        parser_names.append("pdfplumber")
                        for row_number, row in enumerate(table, start=1):
                            for column_number, value in enumerate(row, start=1):
                                cell_text = str(value or "").strip()
                                if not cell_text:
                                    continue
                                location = {"page": page_number, "table": table_number, "row": row_number, "column": column_number}
                                evidence.append(_evidence(kind="pdf_table_cell", location=location, text=cell_text, confidence=0.88))
                                detected.extend(_detect_enterprises(cell_text))
        except Exception as exc:
            issues.append(_issue("table_extraction_failed", "warning", f"PDF表格提取未完成：{type(exc).__name__}。"))
        external_metadata: dict[str, Any] = {"available": False, "provider": None, "model": None, "prompt_version": None}
        if native_text_pages == 0:
            from .external_multimodal import ExternalMultimodalClient, ExternalMultimodalError, render_pdf_pages

            client = external_client or ExternalMultimodalClient.from_environment()
            external_metadata = client.capability()
            if external_metadata["available"]:
                try:
                    parsed = client.parse_pdf_pages(render_pdf_pages(path), source_label=path.name)
                    evidence.extend(parsed.get("evidence", []))
                    detected.extend(parsed.get("detected_enterprise_codes", []))
                    parser_names.append("external_multimodal")
                    external_metadata["response_id"] = parsed.get("response_id")
                    if parsed.get("notes"):
                        metadata_notes = parsed["notes"]
                    else:
                        metadata_notes = []
                except ExternalMultimodalError as exc:
                    issues.append(_issue("external_api_failed", "warning", f"外部扫描PDF解析失败：{exc}"))
                    metadata_notes = []
            else:
                issues.append(_issue("external_api_unavailable", "warning", external_metadata["reason"]))
                metadata_notes = []
            ocr = _ocr_capability()
            if not evidence and not ocr["available"]:
                issues.append(_issue("ocr_unavailable", "warning", ocr["reason"]))
            elif not evidence and ocr["available"]:
                issues.append(_issue("ocr_not_run", "warning", "本地OCR能力已探测，但扫描PDF优先使用外部多模态解析；当前未自动调用本地OCR。"))
        else:
            metadata_notes = []
        metadata = {
            "page_count": page_count,
            "native_text_pages": native_text_pages,
            "table_count": table_count,
            "ocr": _ocr_capability(),
            "external_multimodal": external_metadata,
            "model_notes": metadata_notes,
        }
        return _finalize(
            file_type="pdf",
            parser="+".join(dict.fromkeys(parser_names)) or "pypdf",
            evidence=evidence,
            issues=issues,
            detected_enterprise_codes=detected,
            expected_enterprise_code=expected_enterprise_code,
            metadata=metadata,
        )
    except Exception as exc:
        return _finalize(
            file_type="pdf",
            parser="pypdf",
            evidence=[],
            issues=[_issue("parse_failed", "error", f"PDF无法解析：{type(exc).__name__}: {exc}")],
            detected_enterprise_codes=[],
            expected_enterprise_code=expected_enterprise_code,
            metadata={"page_count": page_count, "ocr": _ocr_capability()},
        )


def _parse_docx(path: Path, expected_enterprise_code: str | None) -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    issues: list[dict[str, Any]] = []
    detected: list[str] = []
    table_count = 0
    image_count = 0
    try:
        from docx import Document

        document = Document(str(path))
        for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
            text = paragraph.text.strip()
            if not text:
                continue
            confidence = 0.95 if paragraph.style and paragraph.style.name else 0.90
            evidence.append(
                _evidence(
                    kind="docx_paragraph",
                    location={"paragraph": paragraph_number, "style": paragraph.style.name if paragraph.style else None},
                    text=text,
                    confidence=confidence,
                )
            )
            detected.extend(_detect_enterprises(text))
        for table_number, table in enumerate(document.tables, start=1):
            table_count += 1
            for row_number, row in enumerate(table.rows, start=1):
                for column_number, cell in enumerate(row.cells, start=1):
                    text = cell.text.strip()
                    if not text:
                        continue
                    evidence.append(
                        _evidence(
                            kind="docx_table_cell",
                            location={"table": table_number, "row": row_number, "column": column_number},
                            text=text,
                            confidence=0.94,
                        )
                    )
                    detected.extend(_detect_enterprises(text))
        image_count = len(getattr(document, "inline_shapes", []))
        if image_count:
            issues.append(_issue("embedded_images_not_ocr", "warning", f"DOCX包含{image_count}个嵌入图片，本轮仅保留引用数量，未对图片执行OCR。"))
        if not evidence:
            issues.append(_issue("empty_document", "error", "DOCX没有可提取的段落或表格内容。"))
        return _finalize(
            file_type="docx",
            parser="python-docx",
            evidence=evidence,
            issues=issues,
            detected_enterprise_codes=detected,
            expected_enterprise_code=expected_enterprise_code,
            metadata={"paragraph_count": len(document.paragraphs), "table_count": table_count, "embedded_image_count": image_count},
        )
    except Exception as exc:
        return _finalize(
            file_type="docx",
            parser="python-docx",
            evidence=[],
            issues=[_issue("parse_failed", "error", f"DOCX无法解析：{type(exc).__name__}: {exc}")],
            detected_enterprise_codes=[],
            expected_enterprise_code=expected_enterprise_code,
            metadata={},
        )


def _parse_image(
    path: Path,
    expected_enterprise_code: str | None,
    external_client: Any | None = None,
) -> dict[str, Any]:
    try:
        from PIL import Image

        with Image.open(path) as image:
            image_format = image.format
            width, height = image.size
            mode = image.mode
            image.load()
            ocr_text, confidence, ocr = _ocr_image(image.copy())
        evidence = []
        detected: list[str] = []
        issues: list[dict[str, Any]] = []
        from .external_multimodal import ExternalMultimodalClient, ExternalMultimodalError

        client = external_client or ExternalMultimodalClient.from_environment()
        external_metadata = client.capability()
        external_notes: list[str] = []
        if external_metadata["available"]:
            try:
                external = client.parse_image(path.read_bytes(), mime_type=f"image/{str(image_format or 'png').lower()}", source_label=path.name)
                evidence.extend(external.get("evidence", []))
                detected.extend(external.get("detected_enterprise_codes", []))
                external_notes = external.get("notes", [])
                external_metadata["response_id"] = external.get("response_id")
            except ExternalMultimodalError as exc:
                issues.append(_issue("external_api_failed", "warning", f"外部图片解析失败：{exc}"))
        if not evidence and ocr_text:
            evidence.append(
                _evidence(
                    kind="image_ocr",
                    location={"page": 1, "coordinates": None},
                    text=ocr_text,
                    confidence=confidence,
                )
            )
            detected = _detect_enterprises(ocr_text)
        elif not evidence and not ocr["available"]:
            issues.append(_issue("ocr_unavailable", "warning", ocr["reason"]))
        elif not evidence:
            issues.append(_issue("ocr_no_text", "warning", "OCR未提取到可用文本，需要人工复核。"))
        result = _finalize(
            file_type="image",
            parser="external_multimodal" if evidence and external_metadata["available"] else "pytesseract" if ocr.get("provider") == "pytesseract" else "image-validation-only",
            evidence=evidence,
            issues=issues,
            detected_enterprise_codes=detected,
            expected_enterprise_code=expected_enterprise_code,
            metadata={
                "format": image_format,
                "width": width,
                "height": height,
                "mode": mode,
                "ocr": ocr,
                "external_multimodal": external_metadata,
                "model_notes": external_notes,
            },
        )
        if not evidence and result["status"] == "passed":
            result["status"] = "needs_review"
            result["merge_allowed"] = False
            result["fact_eligible"] = False
        return result
    except Exception as exc:
        return _finalize(
            file_type="image",
            parser="Pillow",
            evidence=[],
            issues=[_issue("parse_failed", "error", f"图片无法读取：{type(exc).__name__}: {exc}")],
            detected_enterprise_codes=[],
            expected_enterprise_code=expected_enterprise_code,
            metadata={},
        )


def validate_file_bytes(file_type: str, content: bytes) -> str | None:
    if not content:
        return "上传文件为空"
    if file_type == "pdf" and not content.startswith(b"%PDF"):
        return "文件扩展名为PDF，但文件头不符合PDF格式"
    if file_type in {"xlsx", "docx"}:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as archive:
                names = set(archive.namelist())
                required = {"xl/workbook.xml"} if file_type == "xlsx" else {"[Content_Types].xml", "word/document.xml"}
                if not required.issubset(names):
                    return f"文件扩展名为{file_type.upper()}，但不是可识别的Office文件"
        except (OSError, zipfile.BadZipFile):
            return f"文件扩展名为{file_type.upper()}，但Office压缩包已损坏"
    if file_type == "image":
        try:
            from PIL import Image

            with Image.open(io.BytesIO(content)) as image:
                image.verify()
        except Exception as exc:
            return f"图片文件损坏或无法读取：{type(exc).__name__}"
    return None


def parse_file(
    path: Path,
    filename: str,
    expected_enterprise_code: str | None,
    external_client: Any | None = None,
) -> dict[str, Any]:
    file_type = classify_file(filename)
    if file_type is None:
        return _finalize(
            file_type="unknown",
            parser="none",
            evidence=[],
            issues=[_issue("unsupported_extension", "error", f"不支持的文件扩展名：{Path(filename).suffix or '无'}")],
            detected_enterprise_codes=[],
            expected_enterprise_code=expected_enterprise_code,
            metadata={},
        )
    content = path.read_bytes()
    magic_error = validate_file_bytes(file_type, content)
    if magic_error:
        return _finalize(
            file_type=file_type,
            parser="validation-only",
            evidence=[],
            issues=[_issue("invalid_file", "error", magic_error)],
            detected_enterprise_codes=[],
            expected_enterprise_code=expected_enterprise_code,
            metadata={"sha256": sha256_bytes(content), "file_size": len(content)},
        )
    if file_type == "pdf":
        return _parse_pdf(path, expected_enterprise_code, external_client)
    if file_type == "docx":
        return _parse_docx(path, expected_enterprise_code)
    if file_type == "image":
        return _parse_image(path, expected_enterprise_code, external_client)
    return _finalize(
        file_type="xlsx",
        parser="m1-workbook-link",
        evidence=[],
        issues=[_issue("xlsx_requires_m1_batch", "warning", "XLSX事实仍由M1五表校验和批次快照管理，本解析器只负责识别文件类型。")],
        detected_enterprise_codes=[],
        expected_enterprise_code=expected_enterprise_code,
        metadata={"sha256": sha256_bytes(content), "file_size": len(content)},
    )
