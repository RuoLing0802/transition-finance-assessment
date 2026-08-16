from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from PIL import Image
from reportlab.pdfgen import canvas

from app import main
from app.domain_store import DomainStore
from app.parsers.external_multimodal import ExternalMultimodalClient
from app.parsers.multimodal import parse_file as parse_file_with_client
from app.store import BatchStore
from tests.test_m1 import make_workbook, upload


@pytest.fixture()
def multimodal_client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    monkeypatch.setattr(main, "store", BatchStore(tmp_path / "m1-runtime"))
    monkeypatch.setattr(main, "domain_store", DomainStore(tmp_path / "application-data"))
    return TestClient(main.app)


def _run(client: TestClient, workbook: bytes | None = None, code: str = "TFTEST01") -> dict:
    workspace = client.post("/api/v1/workspaces", json={"name": "D阶段多模态测试"}).json()
    workbook = workbook or make_workbook()
    summary = upload(client, workbook)
    source = client.post("/api/v1/source-batches", json={"batch_id": summary["batch_id"]}).json()
    response = client.post(
        f"/api/v1/workspaces/{workspace['workspace_id']}/runs",
        json={
            "enterprise_code": code,
            "source_batch_id": source["source_batch_id"],
            "run_name": f"{code}附件测试",
            "model_config": {"mode": "offline", "provider": "none"},
            "basic_info_index": {"企业代号": code, "索引来源": "基本信息"},
        },
    )
    assert response.status_code == 200, response.text
    return {"run": response.json(), "workbook": workbook, "source": source}


def _pdf(text: str) -> bytes:
    output = BytesIO()
    document = canvas.Canvas(output)
    document.drawString(72, 760, text)
    document.save()
    return output.getvalue()


def _docx(paragraph: str, table_value: str | None = None) -> bytes:
    document = Document()
    document.add_heading("企业补充材料", level=1)
    document.add_paragraph(paragraph)
    if table_value is not None:
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "企业代号"
        table.cell(0, 1).text = table_value
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _image() -> bytes:
    output = BytesIO()
    Image.new("RGB", (120, 60), color="white").save(output, format="PNG")
    return output.getvalue()


def _attach(client: TestClient, run_id: str, filename: str, content: bytes, mime: str) -> dict:
    response = client.post(
        f"/api/v1/assessment-runs/{run_id}/attachments",
        files={"file": (filename, content, mime)},
    )
    assert response.status_code == 200, response.text
    return response.json()


def test_pdf_native_text_has_page_evidence_and_passes(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    result = _attach(multimodal_client, context["run"]["assessment_run_id"], "evidence.pdf", _pdf("企业代号：TFTEST01\n能耗说明"), "application/pdf")
    assert result["parse"]["status"] == "passed"
    assert result["parse"]["merge_allowed"] is True
    assert any(item["location"]["page"] == 1 for item in result["parse"]["evidence"])
    assert result["attachment"]["relative_path"].startswith("attachments/")


def test_docx_paragraph_and_table_locations_are_retained(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    result = _attach(
        multimodal_client,
        context["run"]["assessment_run_id"],
        "evidence.docx",
        _docx("当前企业代号：TFTEST01", "TFTEST01"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert result["parse"]["status"] == "passed"
    locations = [item["location"] for item in result["parse"]["evidence"]]
    assert any("paragraph" in location for location in locations)
    assert any(location.get("table") == 1 and location.get("row") == 1 for location in locations)


def test_missing_enterprise_code_is_low_confidence_and_needs_review(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    result = _attach(multimodal_client, context["run"]["assessment_run_id"], "missing-code.pdf", _pdf("能源补充说明，但没有企业标识"), "application/pdf")
    assert result["parse"]["status"] == "needs_review"
    assert result["parse"]["merge_allowed"] is False
    assert "low_confidence" in {issue["code"] for issue in result["parse"]["issues"]}


def test_enterprise_conflict_blocks_merge_without_cross_run_write(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    run_id = context["run"]["assessment_run_id"]
    result = _attach(multimodal_client, run_id, "other-enterprise.docx", _docx("企业代号：TFTEST02"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    assert result["parse"]["status"] == "blocked_conflict"
    assert result["merge_blocked"] is True
    assert any(issue["code"] == "enterprise_conflict" for issue in result["parse"]["issues"])
    assert multimodal_client.get(f"/api/v1/assessment-runs/{run_id}/attachments").json()["attachments"][0]["assessment_run_id"] == run_id


def test_failed_duplicate_and_xlsx_boundary_are_readable(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    run_id = context["run"]["assessment_run_id"]
    failed = _attach(multimodal_client, run_id, "damaged.pdf", b"not a pdf", "application/pdf")
    assert failed["parse"]["status"] == "failed"
    mime_mismatch = _attach(multimodal_client, run_id, "mime-mismatch.pdf", _pdf("企业代号：TFTEST01"), "text/plain")
    assert mime_mismatch["parse"]["status"] == "needs_review"
    assert any(issue["code"] == "mime_mismatch" for issue in mime_mismatch["parse"]["issues"])
    pdf_content = _pdf("企业代号：TFTEST01")
    duplicate = _attach(multimodal_client, run_id, "evidence.pdf", pdf_content, "application/pdf")
    duplicate_again = _attach(multimodal_client, run_id, "evidence-copy.pdf", pdf_content, "application/pdf")
    assert duplicate_again["attachment"]["reused"] is True
    assert duplicate["attachment"]["sha256"] == duplicate_again["attachment"]["sha256"]
    xlsx = _attach(multimodal_client, run_id, "bound.xlsx", context["workbook"], "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    assert xlsx["parse"]["status"] == "passed"
    unsupported = multimodal_client.post(f"/api/v1/assessment-runs/{run_id}/attachments", files={"file": ("note.txt", b"text", "text/plain")})
    assert unsupported.status_code == 400


def test_image_validation_is_explicit_when_ocr_is_unavailable(multimodal_client: TestClient) -> None:
    context = _run(multimodal_client)
    result = _attach(multimodal_client, context["run"]["assessment_run_id"], "scan.png", _image(), "image/png")
    assert result["attachment"]["file_type"] == "image"
    assert result["parse"]["status"] == "needs_review"
    assert result["parse"]["merge_allowed"] is False
    assert any(issue["code"] in {"ocr_unavailable", "ocr_no_text"} for issue in result["parse"]["issues"])


def test_parser_capability_endpoint_is_truthful(multimodal_client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv("TRANSITION_FINANCE_MULTIMODAL_API_BASE_URL", "https://provider.example/v1")
    monkeypatch.setenv("TRANSITION_FINANCE_MULTIMODAL_API_KEY", "do-not-return")
    monkeypatch.setenv("TRANSITION_FINANCE_MULTIMODAL_MODEL", "vision-test-model")
    body = multimodal_client.get("/api/v1/parsers/capabilities")
    assert body.status_code == 200
    payload = body.json()
    assert payload["file_types"]["pdf"]["available"] is True
    assert payload["file_types"]["docx"]["available"] is True
    assert payload["file_types"]["image"]["available"] is True
    assert "api_key" not in str(payload)
    assert "do-not-return" not in str(payload)


def test_attachment_api_uses_external_parser_and_persists_provenance(multimodal_client: TestClient, monkeypatch: pytest.MonkeyPatch) -> None:
    context = _run(multimodal_client)
    calls: list[tuple[str, dict, dict]] = []

    def transport(url: str, headers: dict[str, str], payload: dict, timeout: float) -> dict:
        calls.append((url, headers, payload))
        return {
            "id": "api-mock-response",
            "choices": [{
                "message": {
                    "content": '{"items":[{"text":"企业代号：TFTEST01","location":{"page":1},"confidence":0.93}],"enterprise_codes":["TFTEST01"]}',
                }
            }],
        }

    client = ExternalMultimodalClient(
        base_url="https://provider.example/v1",
        api_key="secret",
        model="vision-test-model",
        transport=transport,
    )

    def parse_with_mock(path: Path, filename: str, expected: str | None) -> dict:
        return parse_file_with_client(path, filename, expected, external_client=client)

    monkeypatch.setattr(main, "parse_file", parse_with_mock)
    result = _attach(multimodal_client, context["run"]["assessment_run_id"], "api-evidence.png", _image(), "image/png")
    assert result["parse"]["status"] == "passed"
    assert result["parse"]["metadata"]["external_multimodal"]["model"] == "vision-test-model"
    assert result["parse"]["metadata"]["external_multimodal"]["response_id"] == "api-mock-response"
    assert calls


def test_attachment_api_routes_deepseek_upload_through_vision_model(
    multimodal_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    context = _run(multimodal_client)
    client = ExternalMultimodalClient(
        base_url="https://provider.example/v1",
        api_key="secret",
        model="qwen3.6-flash",
        transport=lambda *_args: {
            "choices": [{
                "message": {
                    "content": '{"items":[{"text":"企业代号：TFTEST01","location":{"page":1},"confidence":0.93}],"enterprise_codes":["TFTEST01"]}'
                }
            }]
        },
    )

    def fake_from_environment(cls, *, model_override: str | None = None):
        assert model_override == "qwen3.6-flash"
        return client

    monkeypatch.setattr(ExternalMultimodalClient, "from_environment", classmethod(fake_from_environment))

    def parse_with_client(path: Path, filename: str, expected: str | None, external_client=None) -> dict:
        return parse_file_with_client(path, filename, expected, external_client=external_client)

    monkeypatch.setattr(main, "parse_file", parse_with_client)
    response = multimodal_client.post(
        f"/api/v1/assessment-runs/{context['run']['assessment_run_id']}/attachments",
        data={"session_model_id": "deepseek-v4-pro"},
        files={"file": ("deepseek-upload.png", _image(), "image/png")},
    )
    assert response.status_code == 200, response.text
    routing = response.json()["model_routing"]
    assert routing["requested_model_id"] == "deepseek-v4-pro"
    assert routing["vision_model_id"] == "qwen3.6-flash"
    assert routing["return_to_model_id"] == "deepseek-v4-pro"
    assert routing["switched"] is True
